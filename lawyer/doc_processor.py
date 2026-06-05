"""Загрузка и чанкинг документов (PDF, DOCX, TXT).

Извлечение текста и OCR — только пакеты Python (pip), без внешних программ
(Tesseract, LibreOffice, MS Word, Poppler).
"""

import gc
import html as html_module
import logging
import re
import threading
import uuid
from pathlib import Path
from typing import Any, Callable

from config import (
    LAWYER_CHUNK_OVERLAP,
    LAWYER_CHUNK_SIZE,
    LAWYER_OCR_MAX_PAGES,
    LAWYER_OCR_MAX_SIDE,
    LAWYER_OCR_PAGE_TIMEOUT_SEC,
    LAWYER_OCR_SCALE,
    LAWYER_OCR_SUBPROCESS,
    LAWYER_OCR_TIMEOUT_SEC,
    MAX_LAWYER_PAGES,
)
from lawyer.text_encoding import (
    decode_text_file,
    repair_citation_text,
    repair_text,
    text_quality_score,
)

logger = logging.getLogger(__name__)

CHUNK_SIZE = max(400, LAWYER_CHUNK_SIZE)
CHUNK_OVERLAP = max(50, min(LAWYER_CHUNK_OVERLAP, CHUNK_SIZE // 2))
# Оценка номера страницы для DOCX / «сплошного» текста PDF (символов на страницу A4)
CHARS_PER_PAGE_ESTIMATE = 2400

_last_pdf_hints: list[str] = []

_PYMUPDF_MISSING: bool | None = None
_DOCX_MISSING: bool | None = None
_rapidocr_engine: Any = None

_DOCX_INSTALL_HINT = (
    "Не установлен python-docx (импорт docx). На сервере: "
    "source venv/bin/activate && pip install python-docx "
    "или pip install -r requirements.txt"
)


def pymupdf_available() -> bool:
    global _PYMUPDF_MISSING
    if _PYMUPDF_MISSING is None:
        try:
            import fitz  # noqa: F401
            _PYMUPDF_MISSING = False
        except ImportError:
            _PYMUPDF_MISSING = True
    return not _PYMUPDF_MISSING


def docx_available() -> bool:
    global _DOCX_MISSING
    if _DOCX_MISSING is None:
        try:
            import docx  # noqa: F401
            _DOCX_MISSING = False
        except ImportError:
            _DOCX_MISSING = True
    return not _DOCX_MISSING


def _chars_in_pages(pages: list[dict[str, Any]]) -> int:
    return sum(len(p.get("text", "")) for p in pages)


def _pages_from_full_text(full: str) -> list[dict[str, Any]]:
    if not full.strip():
        return []
    parts = full.split("\x0c") if "\x0c" in full else [full]
    pages: list[dict[str, Any]] = []
    for i, part in enumerate(parts):
        text = _clean_text(part)
        if text:
            pages.append({"page": i + 1, "text": text})
    return pages


def _pdf_security_hint(path: Path) -> str | None:
    """Проверка пароля и запрета копирования текста."""
    try:
        import fitz
    except ImportError:
        return None

    try:
        with fitz.open(path) as doc:
            if doc.needs_pass and not doc.authenticate(""):
                return "PDF защищён паролем — сохраните копию без пароля или загрузите DOCX."
            perms = doc.permissions
            if perms != -1:
                copy_perm = getattr(fitz, "PDF_PERM_COPY", 16)
                if not (perms & copy_perm):
                    return (
                        "В PDF запрещено извлечение текста (защита). "
                        "В Word: Файл → Сохранить как → PDF (новый файл) или загрузите DOCX."
                    )
    except Exception:
        pass
    return None


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = repair_citation_text(text.strip())
    return text


def _read_pdf_pypdf(path: Path) -> list[dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path), strict=False)
    if reader.is_encrypted:
        result = reader.decrypt("")
        if result == 0:
            raise ValueError(
                "PDF защищён паролем. Сохраните копию без пароля или загрузите DOCX/TXT."
            )

    if len(reader.pages) > MAX_LAWYER_PAGES:
        raise ValueError(f"PDF превышает лимит {MAX_LAWYER_PAGES} страниц")

    pages = []
    for i, page in enumerate(reader.pages):
        text = ""
        for use_layout in (True, False):
            try:
                if use_layout:
                    text = page.extract_text(extraction_mode="layout") or ""
                else:
                    text = page.extract_text() or ""
            except TypeError:
                text = page.extract_text() or ""
            if text.strip():
                break

        text = _clean_text(text)
        if text:
            pages.append({"page": i + 1, "text": text})
    return pages


def _read_pdf_pdfplumber(path: Path) -> list[dict[str, Any]]:
    """Запасной движок (pdfminer) — часто читает PDF, где pypdf пустой."""
    try:
        import pdfplumber
    except ImportError:
        return []

    pages: list[dict[str, Any]] = []
    try:
        with pdfplumber.open(path) as pdf:
            if len(pdf.pages) > MAX_LAWYER_PAGES:
                raise ValueError(f"PDF превышает лимит {MAX_LAWYER_PAGES} страниц")
            for i, page in enumerate(pdf.pages):
                text = ""
                for kwargs in ({"layout": True}, {}):
                    try:
                        text = _clean_text(page.extract_text(**kwargs) or "")
                    except TypeError:
                        text = _clean_text(page.extract_text() or "")
                    if text:
                        break
                if text:
                    pages.append({"page": i + 1, "text": text})
    except ValueError:
        raise
    except Exception as e:
        logger.warning("pdfplumber не смог прочитать %s: %s", path.name, e)
        return []

    return pages


def _read_pdf_pdfium(path: Path) -> list[dict[str, Any]]:
    """PDFium — часто читает корпоративные PDF из Word, где pypdf/PyMuPDF пустые."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return []

    pages: list[dict[str, Any]] = []
    pdf = None
    try:
        pdf = pdfium.PdfDocument(str(path))
        n = len(pdf)
        if n > MAX_LAWYER_PAGES:
            raise ValueError(f"PDF превышает лимит {MAX_LAWYER_PAGES} страниц")
        for i in range(n):
            page = pdf[i]
            textpage = page.get_textpage()
            try:
                text = _clean_text(textpage.get_text_range() or "")
            finally:
                textpage.close()
                page.close()
            if text:
                pages.append({"page": i + 1, "text": text})
    except ValueError:
        raise
    except Exception as e:
        logger.warning("pypdfium2 не смог прочитать %s: %s", path.name, e)
        return []
    finally:
        if pdf is not None:
            pdf.close()
    return pages


def _html_to_plain(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
    raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?is)</p>", "\n", raw)
    return _clean_text(html_module.unescape(re.sub(r"<[^>]+>", " ", raw)))


def _fitz_page_text(page: Any) -> str:
    """Несколько способов извлечения текста через PyMuPDF."""
    import fitz

    candidates: list[str] = []
    text_flags: list[int] = [0]
    if hasattr(fitz, "TEXTFLAGS_TEXT"):
        text_flags.insert(0, fitz.TEXTFLAGS_TEXT)
    combo = getattr(fitz, "TEXT_PRESERVE_LIGATURES", 0) | getattr(
        fitz, "TEXT_PRESERVE_WHITESPACE", 0
    )
    if combo:
        text_flags.append(combo)
    for flags in text_flags:
        try:
            raw = page.get_text("text", sort=True, flags=flags) or ""
        except TypeError:
            try:
                raw = page.get_text("text", sort=True) or ""
            except Exception:
                raw = ""
        except Exception:
            raw = ""
        if raw.strip():
            candidates.append(raw)

    for mode in ("blocks", "words", "html", "xhtml"):
        try:
            raw = page.get_text(mode, sort=True) or ""
        except TypeError:
            raw = page.get_text(mode) or ""
        except Exception:
            raw = ""
        if not isinstance(raw, str) or not raw.strip():
            continue
        if mode in ("html", "xhtml"):
            raw = _html_to_plain(raw)
        if raw.strip():
            candidates.append(raw)

    try:
        data = page.get_text("dict", sort=True) or page.get_text("dict") or {}
        span_parts: list[str] = []
        for block in data.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    t = span.get("text", "")
                    if t:
                        span_parts.append(t)
        if span_parts:
            candidates.append("\n".join(span_parts))
    except Exception:
        pass

    if not candidates:
        return ""

    return _clean_text(max(candidates, key=text_quality_score))


def _read_pdf_pymupdf(path: Path) -> list[dict[str, Any]]:
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF (pymupdf) не установлен — pip install pymupdf")
        return []

    pages: list[dict[str, Any]] = []
    try:
        with fitz.open(path) as doc:
            if doc.needs_pass:
                if not doc.authenticate(""):
                    raise ValueError(
                        "PDF защищён паролем. Сохраните копию без пароля или загрузите DOCX/TXT."
                    )
            if len(doc) > MAX_LAWYER_PAGES:
                raise ValueError(f"PDF превышает лимит {MAX_LAWYER_PAGES} страниц")

            for i in range(len(doc)):
                text = _fitz_page_text(doc[i])
                if text:
                    pages.append({"page": i + 1, "text": text})
    except ValueError:
        raise
    except Exception as e:
        logger.warning("PyMuPDF не смог прочитать %s: %s", path.name, e)
        return []

    return pages


def _read_pdf_pymupdf_repair(path: Path) -> list[dict[str, Any]]:
    """Пересборка PDF через PyMuPDF — иногда помогает при битой структуре."""
    try:
        import fitz
    except ImportError:
        return []

    try:
        with fitz.open(path) as src:
            if src.needs_pass and not src.authenticate(""):
                return []
            repaired = fitz.open()
            repaired.insert_pdf(src)
            pdf_bytes = repaired.tobytes(deflate=True)
            repaired.close()
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            if len(doc) > MAX_LAWYER_PAGES:
                raise ValueError(f"PDF превышает лимит {MAX_LAWYER_PAGES} страниц")
            pages: list[dict[str, Any]] = []
            for i in range(len(doc)):
                text = _fitz_page_text(doc[i])
                if text:
                    pages.append({"page": i + 1, "text": text})
            return pages
    except ValueError:
        raise
    except Exception as e:
        logger.warning("PyMuPDF repair не помог для %s: %s", path.name, e)
        return []


def _read_pdf_pdfminer(path: Path) -> list[dict[str, Any]]:
    """Прямой pdfminer — другие параметры разметки, чем у pdfplumber."""
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.layout import LAParams
    except ImportError:
        return []

    try:
        laparams = LAParams(all_texts=True, line_margin=0.12, word_margin=0.1, char_margin=1.0)
        full = extract_text(str(path), laparams=laparams) or ""
    except Exception as e:
        logger.warning("pdfminer не смог прочитать %s: %s", path.name, e)
        return []

    if not full.strip():
        return []

    parts = full.split("\x0c") if "\x0c" in full else [full]
    pages: list[dict[str, Any]] = []
    for i, part in enumerate(parts):
        text = _clean_text(part)
        if text:
            pages.append({"page": i + 1, "text": text})
    return pages


def _page_pixmap_rgb(page: Any, scale: float = 2.0, max_side: int = 2400) -> Any:
    """Растеризация страницы в RGB (без сюрпризов с n/stride)."""
    import fitz

    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False, colorspace=fitz.csRGB)
    if max(pix.width, pix.height) > max_side:
        shrink = max_side / max(pix.width, pix.height)
        mat = fitz.Matrix(scale * shrink, scale * shrink)
        pix = page.get_pixmap(matrix=mat, alpha=False, colorspace=fitz.csRGB)
    return pix


def _pixmap_to_numpy(pix: Any) -> Any:
    """Pixmap PyMuPDF → RGB numpy (H, W, 3), с учётом stride."""
    import numpy as np
    from PIL import Image

    w, h, n = pix.width, pix.height, pix.n
    if w <= 0 or h <= 0:
        raise ValueError(f"некорректный размер pixmap: {w}x{h}")

    try:
        if n not in (1, 3, 4):
            raise ValueError(f"неожиданное число каналов: {n}")

        row_bytes = w * n
        if pix.stride == row_bytes:
            arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(h, w, n)
        else:
            arr = np.zeros((h, w, n), dtype=np.uint8)
            for y in range(h):
                off = y * pix.stride
                end = off + row_bytes
                if end > len(pix.samples):
                    raise IndexError(
                        f"stride={pix.stride}, samples={len(pix.samples)}, row={y}"
                    )
                arr[y] = np.frombuffer(
                    pix.samples[off:end], dtype=np.uint8
                ).reshape(w, n)

        if n == 1:
            rgb = np.stack([arr[:, :, 0]] * 3, axis=-1)
        elif n == 4:
            rgb = arr[:, :, :3].copy()
        else:
            rgb = arr

        return np.ascontiguousarray(rgb)
    except (ValueError, IndexError, BufferError) as e:
        logger.debug("Pixmap через PIL (fallback): %s", e)
        img = Image.frombytes("RGB", (w, h), pix.samples)
        return np.ascontiguousarray(img)


def _get_rapidocr_engine() -> Any:
    global _rapidocr_engine
    if _rapidocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR

        logger.info("RapidOCR: загрузка модели (первый раз может занять 10–30 с)...")
        _rapidocr_engine = RapidOCR()
    return _rapidocr_engine


_OCR_OOM_HINT = (
    "OCR прерван (нехватает RAM на сервере). В .env: LAWYER_OCR_SCALE=1.0, "
    "LAWYER_OCR_MAX_SIDE=1200, добавьте swap 2 ГБ или загрузите DOCX."
)

_OcrProgressCallback = Callable[
    [list[dict[str, Any]], int, int, str | None], None
]


def _ocr_result_to_lines(result: Any) -> list[str]:
    """Безопасный разбор ответа RapidOCR."""
    lines: list[str] = []
    if not result:
        return lines
    for item in result:
        if not item:
            continue
        try:
            if isinstance(item, (list, tuple)) and len(item) >= 2 and item[1]:
                lines.append(str(item[1]))
            elif isinstance(item, dict) and item.get("text"):
                lines.append(str(item["text"]))
        except (IndexError, TypeError, KeyError):
            continue
    return lines


def _run_rapidocr_engine(engine: Any, img: Any, timeout_sec: int) -> Any:
    """RapidOCR с таймаутом на страницу (защита от зависания на Windows)."""
    box: dict[str, Any] = {"result": None, "error": None}

    def _target() -> None:
        try:
            out = engine(img)
            box["result"] = out[0] if isinstance(out, tuple) else out
        except Exception as e:
            box["error"] = e

    thread = threading.Thread(target=_target, daemon=True)
    thread.start()
    thread.join(timeout=max(10, timeout_sec))
    if thread.is_alive():
        raise TimeoutError(f"OCR страницы > {timeout_sec} с")
    if box["error"] is not None:
        raise box["error"]
    return box["result"]


def _ocr_page_attempts(scale: float, max_side: int) -> list[tuple[float, int]]:
    """Параметры растеризации: основной и два запасных (меньше — быстрее и стабильнее)."""
    return [
        (scale, max_side),
        (max(0.8, scale * 0.85), max_side),
        (scale, max(800, int(max_side * 0.75))),
    ]


def _rapidocr_worker(path_str: str, out_queue: Any) -> None:
    """Точка входа дочернего процесса OCR."""

    def _progress(
        pages: list[dict[str, Any]], done: int, total: int, err: str | None
    ) -> None:
        out_queue.put(("progress", list(pages), err, done, total))

    try:
        pages, err = _read_pdf_rapidocr_impl(
            Path(path_str),
            progress_callback=_progress,
        )
        out_queue.put(("ok", pages, err))
    except Exception as e:
        out_queue.put(("err", [], str(e)))


def _drain_ocr_queue(
    queue: Any,
    partial_pages: list[dict[str, Any]],
    partial_err: list[str | None],
) -> tuple[str, list[dict[str, Any]], str | None] | None:
    """Забрать из очереди прогресс или финальный результат subprocess."""
    from queue import Empty

    while True:
        try:
            msg = queue.get_nowait()
        except Empty:
            return None
        if not msg:
            continue
        kind = msg[0]
        if kind == "progress":
            partial_pages[:] = msg[1]
            partial_err[0] = msg[2]
        elif kind in ("ok", "err"):
            return (kind, msg[1], msg[2])


def _read_pdf_rapidocr_subprocess(path: Path) -> tuple[list[dict[str, Any]], str | None]:
    """OCR в отдельном процессе: при Killed/OOM основной сервер остаётся жив."""
    import multiprocessing as mp
    import time

    ctx = mp.get_context("spawn")
    queue: Any = ctx.Queue()
    proc = ctx.Process(
        target=_rapidocr_worker,
        args=(str(path.resolve()), queue),
        daemon=True,
    )
    proc.start()
    deadline = time.monotonic() + LAWYER_OCR_TIMEOUT_SEC
    partial_pages: list[dict[str, Any]] = []
    partial_err: list[str | None] = [None]

    while proc.is_alive():
        remaining = deadline - time.monotonic()
        if remaining <= 0:
            break
        proc.join(timeout=min(1.0, remaining))
        done = _drain_ocr_queue(queue, partial_pages, partial_err)
        while done:
            if done[0] == "ok":
                proc.join(timeout=5)
                return done[1], done[2]
            if done[0] == "err":
                if partial_pages:
                    return partial_pages, done[2] or partial_err[0]
                return [], done[2] or "OCR завершился с ошибкой"
            done = _drain_ocr_queue(queue, partial_pages, partial_err)

    while True:
        done = _drain_ocr_queue(queue, partial_pages, partial_err)
        if not done:
            break
        if done[0] == "ok":
            return done[1], done[2]

    if proc.is_alive():
        proc.terminate()
        proc.join(10)
        _drain_ocr_queue(queue, partial_pages, partial_err)
        if partial_pages:
            tail = (
                f" частично распознано {len(partial_pages)} стр. "
                f"(таймаут {LAWYER_OCR_TIMEOUT_SEC} с)"
            )
            err = (partial_err[0] or "") + tail
            logger.warning(
                "RapidOCR: таймаут %d с, сохранено %d стр.",
                LAWYER_OCR_TIMEOUT_SEC,
                len(partial_pages),
            )
            return partial_pages, err.strip() or None
        logger.warning("RapidOCR: таймаут %d с", LAWYER_OCR_TIMEOUT_SEC)
        return [], f"OCR превысил лимит времени ({LAWYER_OCR_TIMEOUT_SEC} с)."

    if proc.exitcode not in (0, None):
        logger.warning("RapidOCR subprocess exitcode=%s", proc.exitcode)
        if partial_pages:
            return partial_pages, partial_err[0] or _OCR_OOM_HINT
        return [], _OCR_OOM_HINT

    try:
        msg = queue.get(timeout=5)
    except Exception:
        if partial_pages:
            return partial_pages, partial_err[0]
        return [], _OCR_OOM_HINT

    if msg[0] == "ok":
        return msg[1], msg[2]
    if msg[0] == "err":
        if partial_pages:
            return partial_pages, str(msg[2]) if msg[2] else partial_err[0]
        return [], str(msg[2]) if msg[2] else "OCR завершился с ошибкой"
    if partial_pages:
        return partial_pages, partial_err[0]
    return [], partial_err[0] or _OCR_OOM_HINT


def _read_pdf_rapidocr(path: Path) -> tuple[list[dict[str, Any]], str | None]:
    if LAWYER_OCR_SUBPROCESS:
        return _read_pdf_rapidocr_subprocess(path)
    return _read_pdf_rapidocr_impl(path)


def _read_pdf_rapidocr_impl(
    path: Path,
    progress_callback: _OcrProgressCallback | None = None,
) -> tuple[list[dict[str, Any]], str | None]:
    """OCR для PDF-сканов (RapidOCR + PyMuPDF, только pip-пакеты)."""
    try:
        import fitz
        import numpy as np
    except ImportError as e:
        return [], f"нет зависимости: {e.name}"

    try:
        engine = _get_rapidocr_engine()
    except ImportError as e:
        logger.warning("RapidOCR ImportError: %s", e)
        err = str(e)
        if "libGL" in err:
            hint = (
                "На сервере установлен opencv-python вместо headless. Выполните: "
                "pip uninstall -y opencv-python && "
                "pip install opencv-python-headless"
            )
        else:
            hint = (
                "pip install rapidocr-onnxruntime onnxruntime opencv-python-headless"
            )
        return [], f"OCR не запущен ({err}). {hint}"
    except Exception as e:
        logger.warning("RapidOCR init failed: %s", e)
        return [], f"RapidOCR не запустился: {e}"

    pages: list[dict[str, Any]] = []
    ocr_error: str | None = None
    scale = max(0.8, min(LAWYER_OCR_SCALE, 3.0))
    max_side = max(800, min(LAWYER_OCR_MAX_SIDE, 3200))
    failed_pages = 0

    try:
        with fitz.open(path) as doc:
            n_pages = len(doc)
            if n_pages == 0:
                return [], "PDF не содержит страниц или файл повреждён при загрузке"
            if n_pages > MAX_LAWYER_PAGES:
                raise ValueError(f"PDF превышает лимит {MAX_LAWYER_PAGES} страниц")

            ocr_limit = (
                min(n_pages, LAWYER_OCR_MAX_PAGES)
                if LAWYER_OCR_MAX_PAGES > 0
                else n_pages
            )
            logger.info(
                "RapidOCR: %d/%d стр., scale=%.2f, max_side=%d — %s",
                ocr_limit,
                n_pages,
                scale,
                max_side,
                path.name,
            )
            for i in range(ocr_limit):
                page_text: str | None = None
                last_err: Exception | None = None
                for attempt_idx, (page_scale, page_max_side) in enumerate(
                    _ocr_page_attempts(scale, max_side)
                ):
                    try:
                        pix = _page_pixmap_rgb(
                            doc[i], scale=page_scale, max_side=page_max_side
                        )
                        img = _pixmap_to_numpy(pix)
                        del pix

                        result = _run_rapidocr_engine(
                            engine,
                            img,
                            LAWYER_OCR_PAGE_TIMEOUT_SEC,
                        )
                        del img

                        lines = _ocr_result_to_lines(result)
                        page_text = _clean_text("\n".join(lines))
                        if page_text:
                            break
                        if attempt_idx == 0 and i == 0:
                            logger.info("RapidOCR: на 1-й странице текст не найден")
                        break
                    except Exception as page_err:
                        last_err = page_err
                        if attempt_idx < len(_ocr_page_attempts(scale, max_side)) - 1:
                            logger.info(
                                "RapidOCR: стр. %d/%d повтор после %s",
                                i + 1,
                                ocr_limit,
                                page_err,
                            )
                            gc.collect()
                            continue
                        failed_pages += 1
                        logger.warning(
                            "RapidOCR: страница %d/%d — %s",
                            i + 1,
                            ocr_limit,
                            page_err,
                        )
                        gc.collect()

                if page_text:
                    pages.append({"page": i + 1, "text": page_text})
                elif last_err is None and not page_text:
                    pass

                if (i + 1) % 3 == 0 or i + 1 == ocr_limit:
                    logger.info("RapidOCR: обработано %d/%d стр.", i + 1, ocr_limit)
                if progress_callback:
                    progress_callback(pages, i + 1, ocr_limit, ocr_error)
            if n_pages > ocr_limit:
                tail = (
                    f" Распознаны только первые {ocr_limit} из {n_pages} стр. "
                    f"(LAWYER_OCR_MAX_PAGES). Для полного текста загрузите DOCX."
                )
                ocr_error = (ocr_error or "") + tail
            if failed_pages and not pages:
                ocr_error = f"OCR не распознал ни одной из {ocr_limit} страниц"
            elif failed_pages:
                ocr_error = f"пропущено страниц: {failed_pages}"
    except ValueError:
        raise
    except Exception as e:
        ocr_error = str(e)
        logger.warning("RapidOCR: %s", e)

    return pages, ocr_error


def _pdf_needs_ocr_only(path: Path) -> bool:
    """Мало текста в слое PDF — нужен OCR (скан / картинки)."""
    try:
        import fitz
    except ImportError:
        return False

    try:
        with fitz.open(path) as doc:
            if not len(doc):
                return False
            for i in range(min(3, len(doc))):
                if len((doc[i].get_text() or "").strip()) > 40:
                    return False
            return True
    except Exception:
        return False


def _rapidocr_available() -> bool:
    try:
        from rapidocr_onnxruntime import RapidOCR  # noqa: F401
        return True
    except ImportError:
        return False


def _diagnose_empty_pdf(path: Path) -> list[str]:
    hints: list[str] = []
    try:
        size = path.stat().st_size
        hints.append(f"размер {size} байт")
    except OSError as e:
        hints.append(f"файл недоступен: {e}")
        return hints

    with open(path, "rb") as f:
        if f.read(5) != b"%PDF-":
            hints.append("нет заголовка %PDF — файл повреждён при загрузке")
            return hints

    try:
        import fitz

        with fitz.open(path) as doc:
            n = len(doc)
            hints.append(f"страниц: {n}")
            pages_with_images = 0
            for i in range(n):
                if doc[i].get_images():
                    pages_with_images += 1
            if n and pages_with_images >= max(1, n * 0.8):
                hints.append(
                    "похоже на скан (картинки) — для OCR нужен rapidocr-onnxruntime"
                )
    except Exception as e:
        hints.append(f"анализ PyMuPDF: {e}")

    if _rapidocr_available():
        hints.append("RapidOCR: установлен")
    else:
        hints.append(
            "RapidOCR: не установлен — pip install rapidocr-onnxruntime opencv-python-headless"
        )

    if not pymupdf_available():
        hints.append("PyMuPDF: не установлен — pip install pymupdf")

    return hints


def _read_pdf(path: Path) -> list[dict[str, Any]]:
    global _last_pdf_hints
    hints: list[str] = []
    _last_pdf_hints = []

    sec = _pdf_security_hint(path)
    if sec:
        hints.append(sec)
        logger.warning("PDF %s: %s", path.name, sec)

    steps: list[tuple[str, Callable[[Path], list[dict[str, Any]]]]] = []
    needs_ocr = _pdf_needs_ocr_only(path)
    if needs_ocr:
        logger.info("PDF без текстового слоя — после извлечения текста будет RapidOCR")

    if pymupdf_available():
        steps.extend([
            ("PyMuPDF", _read_pdf_pymupdf),
        ])
        if not needs_ocr:
            steps.extend([
                ("pypdfium2", _read_pdf_pdfium),
                ("pdfplumber", _read_pdf_pdfplumber),
                ("pdfminer", _read_pdf_pdfminer),
                ("pypdf", _read_pdf_pypdf),
                ("PyMuPDF-repair", _read_pdf_pymupdf_repair),
            ])
        else:
            # Скан: перед тяжёлым OCR пробуем другие извлекатели
            steps.extend([
                ("pypdfium2", _read_pdf_pdfium),
                ("pdfplumber", _read_pdf_pdfplumber),
                ("PyMuPDF-repair", _read_pdf_pymupdf_repair),
            ])
    else:
        hints.append("не установлен pymupdf — pip install pymupdf")
        steps.extend([
            ("pypdfium2", _read_pdf_pdfium),
            ("pdfplumber", _read_pdf_pdfplumber),
            ("pdfminer", _read_pdf_pdfminer),
            ("pypdf", _read_pdf_pypdf),
        ])

    for name, reader in steps:
        logger.info("PDF: пробуем %s — %s", name, path.name)
        try:
            pages = reader(path)
        except ValueError:
            raise
        except Exception as e:
            logger.warning("PDF %s (%s): %s", path.name, name, e)
            pages = []
        logger.info(
            "PDF: %s → %d стр., %d симв.",
            name,
            len(pages),
            _chars_in_pages(pages),
        )
        if pages:
            logger.info("PDF (%s): успех — %s", name, path.name)
            return pages

    logger.info("PDF: пробуем RapidOCR — %s", path.name)
    try:
        import fitz

        with fitz.open(path) as doc:
            logger.info(
                "PDF: %d стр. для OCR (лимит %s), таймаут стр.=%d с — %s",
                len(doc),
                LAWYER_OCR_MAX_PAGES or "все",
                LAWYER_OCR_PAGE_TIMEOUT_SEC,
                path.name,
            )
    except Exception:
        pass
    pages, ocr_err = _read_pdf_rapidocr(path)
    if pages:
        logger.info("PDF (RapidOCR): %d стр. — %s", len(pages), path.name)
        return pages
    if ocr_err:
        logger.warning("PDF RapidOCR не дал текста (%s): %s", path.name, ocr_err)
        hints.append(ocr_err)

    hints.extend(_diagnose_empty_pdf(path))
    _last_pdf_hints = hints
    return []


def _pdf_failure_message(path: Path) -> str:
    if _last_pdf_hints and any("запрещено извлечение" in h for h in _last_pdf_hints):
        return _last_pdf_hints[0]
    if _last_pdf_hints and any("pymupdf" in h.lower() for h in _last_pdf_hints):
        return (
            "Не удалось прочитать PDF: установите зависимости Python:\n"
            "pip install pymupdf pypdfium2 pdfplumber pypdf\n"
            "Перезапустите сервер."
        )
    if _last_pdf_hints and any(
        "скан" in h or "RapidOCR" in h or "OCR" in h for h in _last_pdf_hints
    ):
        detail = "; ".join(_last_pdf_hints[:4])
        return (
            "PDF без текстового слоя (скан). На сервере не удалось распознать текст.\n"
            "Если в логе «Killed» — не хватает RAM: добавьте swap, уменьшите "
            "LAWYER_OCR_SCALE=1.0 и LAWYER_OCR_MAX_SIDE=1200 в .env, либо загрузите DOCX.\n"
            "OCR: pip install rapidocr-onnxruntime onnxruntime opencv-python-headless\n"
            f"Диагностика: {detail}"
        )
    base = (
        "Не удалось извлечь текст из PDF средствами Python. "
        "Загрузите DOCX/TXT или установите полный набор: "
        "pip install -r requirements.txt"
    )
    if _last_pdf_hints:
        return f"{base} Диагностика: {'; '.join(_last_pdf_hints)}."
    return base


def _read_docx(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ModuleNotFoundError as e:
        raise RuntimeError(_DOCX_INSTALL_HINT) from e

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    full_text = _clean_text(full_text)
    if not full_text:
        raise ValueError("DOCX не содержит текста")
    return [{"page": 1, "text": full_text}]


def _read_txt(path: Path) -> list[dict[str, Any]]:
    text = _clean_text(decode_text_file(path))
    if not text:
        raise ValueError("TXT-файл пуст")
    return [{"page": 1, "text": text}]


def _normalize_pdf_pages(path: Path, pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Если PDF многостраничный, а текст слился в одну «страницу» — разбить по листам PyMuPDF."""
    if not pages:
        return pages
    try:
        import fitz
    except ImportError:
        return pages

    try:
        with fitz.open(path) as doc:
            pdf_pages = len(doc)
            if pdf_pages <= 1:
                return pages
            merged_one = len(pages) == 1 and len(pages[0].get("text", "")) > 400
            if not merged_one:
                return pages

            per_page: list[dict[str, Any]] = []
            for i in range(pdf_pages):
                text = _fitz_page_text(doc[i])
                if text:
                    per_page.append({"page": i + 1, "text": text})
            if per_page:
                logger.info(
                    "PDF %s: переразбивка на %d стр. (был один блок текста)",
                    path.name,
                    len(per_page),
                )
                return per_page
    except Exception as e:
        logger.warning("Не удалось переразбить PDF по страницам %s: %s", path.name, e)
    return pages


def load_document(path: Path) -> list[dict[str, Any]]:
    """Чтение документа с разбивкой по страницам."""
    if not path.exists() or path.stat().st_size == 0:
        raise ValueError("Файл пуст или не найден")

    ext = path.suffix.lower()
    if not ext:
        with open(path, "rb") as f:
            if f.read(5) == b"%PDF-":
                ext = ".pdf"
    if ext == ".pdf":
        pages = _read_pdf(path)
        return _normalize_pdf_pages(path, pages)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".txt":
        return _read_txt(path)
    raise ValueError(f"Неподдерживаемый формат: {ext}")


def _page_for_chunk(
    *,
    filename: str,
    pages: list[dict[str, Any]],
    page_num: int,
    file_char_offset: int,
    chunk_start: int,
) -> int:
    """Номер страницы для чанка: из PDF-листа или оценка для DOCX/сплошного текста."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx" or (ext == ".pdf" and len(pages) == 1):
        return max(1, (file_char_offset + chunk_start) // CHARS_PER_PAGE_ESTIMATE + 1)
    return max(1, int(page_num))


def _next_chunk_end(text: str, start: int, size: int) -> int:
    """Конец чанка: по возможности не резать списки и абзацы."""
    hard_end = min(start + size, len(text))
    if hard_end >= len(text):
        return hard_end

    min_end = start + size // 2
    for marker in ("\n\n", "\n•", "\n- ", "\n— "):
        pos = text.rfind(marker, start, hard_end)
        if pos >= min_end:
            return pos + len(marker)

    for match in reversed(list(re.finditer(r"\n\d+[\.)]\s", text[start:hard_end]))):
        end = start + match.start()
        if end >= min_end:
            return end

    nl = text.rfind("\n", start, hard_end)
    if nl >= min_end:
        return nl + 1
    return hard_end


def chunk_text(
    pages: list[dict[str, Any]],
    filename: str,
    file_id: str,
) -> list[dict[str, Any]]:
    """Разбиение текста на чанки с перекрытием."""
    chunks = []
    chunk_idx = 0
    file_char_offset = 0

    for page_data in pages:
        text = page_data["text"]
        page_num = int(page_data.get("page") or 1)
        start = 0
        while start < len(text):
            end = _next_chunk_end(text, start, CHUNK_SIZE)
            chunk_text_str = text[start:end]
            if chunk_text_str.strip():
                cite_page = _page_for_chunk(
                    filename=filename,
                    pages=pages,
                    page_num=page_num,
                    file_char_offset=file_char_offset,
                    chunk_start=start,
                )
                chunks.append({
                    "id": f"{file_id}_{chunk_idx}",
                    "text": repair_citation_text(chunk_text_str.strip()),
                    "metadata": {
                        "file_id": file_id,
                        "filename": repair_text(filename),
                        "page": cite_page,
                        "chunk_index": chunk_idx,
                    },
                })
                chunk_idx += 1
            start += CHUNK_SIZE - CHUNK_OVERLAP
            if start >= len(text):
                break
        file_char_offset += len(text)

    logger.info("Документ %s: %d чанков", filename, len(chunks))
    return chunks


def process_upload(path: Path, filename: str) -> tuple[str, list[dict[str, Any]]]:
    """Полная обработка загруженного файла."""
    file_id = str(uuid.uuid4())[:12]
    pages = load_document(path)
    if not pages:
        if path.suffix.lower() == ".pdf":
            raise ValueError(_pdf_failure_message(path))
        raise ValueError("Документ пуст или не удалось извлечь текст")

    chunks = chunk_text(pages, filename, file_id)
    if not chunks:
        raise ValueError("Текст извлечён, но слишком короткий для индексации")
    return file_id, chunks
